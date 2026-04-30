"""
Script to generate the professional project report as a .docx file.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Colour palette ──────────────────────────────────────────────────────────
NAVY      = RGBColor(0x0D, 0x2B, 0x45)   # heading bg
BLUE      = RGBColor(0x29, 0x80, 0xB9)   # accent / links
DARK_BLUE = RGBColor(0x1A, 0x5F, 0x8A)   # subheading
RED       = RGBColor(0xC0, 0x39, 0x2B)   # danger
GREEN     = RGBColor(0x1A, 0x8A, 0x4A)   # safe
ORANGE    = RGBColor(0xE6, 0x7E, 0x22)   # warning
LIGHT_BG  = RGBColor(0xF4, 0xF7, 0xFA)   # table row alt
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x1A, 0x1A, 0x1A)
GREY      = RGBColor(0x55, 0x55, 0x55)
MID_GREY  = RGBColor(0xCC, 0xCC, 0xCC)


# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    """Apply a solid background colour to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Set borders on a table cell (top/bottom/left/right)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   kwargs.get(side, {}).get("val",   "single"))
        border.set(qn("w:sz"),    kwargs.get(side, {}).get("sz",    "4"))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), kwargs.get(side, {}).get("color", "CCCCCC"))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc: Document, text: str, level: int = 1,
                color: RGBColor = NAVY, space_before: int = 18,
                space_after: int = 6):
    """Add a styled heading paragraph."""
    p = doc.add_heading(text, level=level)
    p.clear()                  # remove default style formatting
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(20)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), f"{BLUE[0]:02X}{BLUE[1]:02X}{BLUE[2]:02X}")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p


def add_body(doc: Document, text: str, indent: float = 0,
             color: RGBColor = BLACK, italic: bool = False,
             bold: bool = False, size: int = 11):
    """Add a body paragraph."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size       = Pt(size)
    run.font.color.rgb  = color
    run.italic          = italic
    run.bold            = bold
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def add_bullet(doc: Document, text: str, level: int = 0,
               color: RGBColor = BLACK, bold_prefix: str = ""):
    """Add a bullet-list item."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.size      = Pt(11)
    r2.font.color.rgb = color
    p.paragraph_format.left_indent   = Inches(0.25 * (level + 1))
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(1)
    return p


def add_numbered(doc: Document, text: str, bold_prefix: str = ""):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    return p


def add_info_table(doc: Document, rows: list[tuple[str, str]],
                   col_widths=(2.0, 4.5)):
    """Render a two-column key-value info table."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (key, val) in enumerate(rows):
        c0, c1 = tbl.rows[i].cells
        c0.width = Inches(col_widths[0])
        c1.width = Inches(col_widths[1])
        c0.text = key
        c1.text = val
        bg = LIGHT_BG if i % 2 == 0 else WHITE
        set_cell_bg(c0, bg); set_cell_bg(c1, bg)
        for cell in (c0, c1):
            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
            cell.paragraphs[0].runs[0].font.color.rgb = BLACK
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c0.paragraphs[0].runs[0].bold = True
        c0.paragraphs[0].runs[0].font.color.rgb = DARK_BLUE
        set_cell_border(c0); set_cell_border(c1)
    doc.add_paragraph()


def add_section_table(doc: Document, headers: list[str],
                      data: list[list[str]]):
    """Render a multi-column data table with a blue header row."""
    tbl = doc.add_table(rows=1 + len(data), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = tbl.rows[0]
    for j, h in enumerate(headers):
        cell = hdr.cells[j]
        cell.text = h
        set_cell_bg(cell, NAVY)
        r = cell.paragraphs[0].runs[0]
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size      = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, **{s: {"val": "single", "sz": "4", "color": "1A5F8A"}
                                  for s in ("top", "left", "bottom", "right")})
    # Data rows
    for i, row_data in enumerate(data):
        row = tbl.rows[i + 1]
        bg  = LIGHT_BG if i % 2 == 0 else WHITE
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            r = cell.paragraphs[0].runs[0]
            r.font.size      = Pt(10)
            r.font.color.rgb = BLACK
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str):
    """Render a monospaced code block with a light-grey background."""
    for line in text.strip().splitlines():
        p = doc.add_paragraph()
        r = p.add_run(line if line.strip() else " ")
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x20, 0x20, 0x80)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F0F4F8")
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)
    doc.add_paragraph()


def add_divider(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


def add_callout(doc: Document, text: str, style: str = "info"):
    """Add a highlighted callout box (info / warning / danger)."""
    colours = {
        "info":    (RGBColor(0xEB, 0xF5, 0xFB), BLUE),
        "warning": (RGBColor(0xFE, 0xF9, 0xE7), ORANGE),
        "danger":  (RGBColor(0xFD, 0xED, 0xEC), RED),
        "success": (RGBColor(0xE9, 0xF7, 0xEF), GREEN),
    }
    bg, fg = colours.get(style, colours["info"])
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size      = Pt(10.5)
    r.font.color.rgb = fg
    r.italic         = True
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{bg[0]:02X}{bg[1]:02X}{bg[2]:02X}")
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p


# ─── Document assembly ────────────────────────────────────────────────────────

def build_report(output_path: str = "Project_Report.docx"):
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width        = Inches(8.5)
    section.page_height       = Inches(11)
    section.left_margin       = Inches(1.0)
    section.right_margin      = Inches(1.0)
    section.top_margin        = Inches(1.0)
    section.bottom_margin     = Inches(1.0)

    # Default body font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ================================================================
    # COVER PAGE
    # ================================================================
    doc.add_paragraph("\n\n\n")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("Email Header Phishing\nInvestigation Tool")
    tr.bold           = True
    tr.font.size      = Pt(32)
    tr.font.color.rgb = NAVY

    doc.add_paragraph()
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("Comprehensive Technical Project Report")
    sr.font.size      = Pt(16)
    sr.font.color.rgb = DARK_BLUE
    sr.italic         = True

    doc.add_paragraph()
    # Decorative horizontal rule
    div_p = doc.add_paragraph()
    div_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_r = div_p.add_run("─" * 60)
    div_r.font.color.rgb = BLUE
    div_r.font.size      = Pt(12)

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_r = meta_p.add_run(
        f"Prepared by: Maryam Inam\n"
        f"Date: {datetime.date(2026, 4, 16).strftime('%B %d, %Y')}\n"
        f"Version: 1.0  |  Classification: Internal\n"
        f"Platform: Windows 11  |  Language: Python 3.10+"
    )
    meta_r.font.size      = Pt(12)
    meta_r.font.color.rgb = GREY

    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS (manual)
    # ================================================================
    add_heading(doc, "Table of Contents", 1)

    toc_items = [
        ("1.", "Executive Summary",                         4),
        ("2.", "Project Overview",                          4),
        ("3.", "Objectives & Scope",                        5),
        ("4.", "System Architecture",                       5),
        ("5.", "Technology Stack",                          6),
        ("6.", "Core Features & Functionality",             7),
        ("7.", "Backend Analysis Engine",                   8),
        ("8.", "Graphical User Interface (GUI)",            10),
        ("9.", "Threat Intelligence Integration",           12),
        ("10.", "Risk Scoring Methodology",                 13),
        ("11.", "Data Flow & Processing Pipeline",          14),
        ("12.", "Algorithms & Technical Implementation",    15),
        ("13.", "Security Considerations",                  17),
        ("14.", "Testing & Sample Data",                    18),
        ("15.", "Installation & Deployment",                18),
        ("16.", "Known Limitations & Future Enhancements",  19),
        ("17.", "Conclusion",                               20),
        ("Appendix A", "File Structure",                    21),
        ("Appendix B", "Risk Score Weight Reference",       21),
        ("Appendix C", "Supported Risky File Extensions",   22),
        ("Appendix D", "Phishing Pattern Categories",       22),
    ]
    for num, title, _ in toc_items:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{num:<6}")
        r1.bold = True
        r1.font.size      = Pt(11)
        r1.font.color.rgb = DARK_BLUE
        r2 = p.add_run(title)
        r2.font.size      = Pt(11)
        r2.font.color.rgb = BLACK
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))

    doc.add_page_break()

    # ================================================================
    # 1. EXECUTIVE SUMMARY
    # ================================================================
    add_heading(doc, "1. Executive Summary", 1)
    add_body(doc,
        "The Email Header Phishing Investigation Tool is a professional-grade, open-source "
        "desktop application developed in Python. It enables cybersecurity analysts, digital "
        "forensic investigators, incident response teams, and security-aware end users to "
        "perform deep forensic analysis of email messages in order to detect, investigate, and "
        "document phishing attacks.")
    add_body(doc,
        "The tool operates by parsing RFC 5322-compliant email files (.eml format), extracting "
        "all relevant headers, tracing the email routing path, validating authentication "
        "mechanisms (SPF, DKIM, DMARC), identifying suspicious URLs and attachments, querying "
        "external threat-intelligence APIs, and producing a comprehensive risk score with an "
        "exportable forensic report.")
    add_body(doc,
        "The application is built on a two-layer architecture: a headless backend analysis "
        "engine (email_forensic_analyzer.py) and a modern CustomTkinter-based graphical "
        "interface (forensic_gui.py). The tool comprises approximately 3,841 lines of Python "
        "code, 10 analysis tabs, 8 analysis phases, and integrates with three external APIs.")

    add_callout(doc,
        "Key Achievement: The tool successfully automates what would otherwise require manual "
        "inspection of hundreds of raw email headers, reducing analysis time from hours to "
        "seconds while providing a professional, exportable forensic report.",
        "success")

    add_heading(doc, "Key Highlights", 2)
    highlights = [
        ("Analysis Engine:", " 3,841 lines of Python across 2 core modules"),
        ("Analysis Phases:", " 8–10 phases covering headers, auth, routing, URLs, attachments, and OSINT"),
        ("External APIs:", " ip-api.com (geolocation), AbuseIPDB (IP reputation), VirusTotal (malware)"),
        ("Threat Detection:", " 50+ regex phishing patterns across 3 categories"),
        ("Risk Scoring:", " Weighted 0–100 score with four threat levels (Low / Medium / High / Critical)"),
        ("GUI Tabs:", " 10 navigational tabs with real-time data population"),
        ("Export Formats:", " HTML reports, JSON IOCs, CSV IOCs, batch summaries"),
        ("Dependencies:", " 5 third-party packages + standard library only"),
    ]
    for bold, normal in highlights:
        add_bullet(doc, normal, bold_prefix=bold)

    doc.add_page_break()

    # ================================================================
    # 2. PROJECT OVERVIEW
    # ================================================================
    add_heading(doc, "2. Project Overview", 1)
    add_info_table(doc, [
        ("Project Name",       "Email Header Phishing Investigation Tool"),
        ("Type",               "Desktop Forensic Investigation Application"),
        ("Primary Language",   "Python 3.10+"),
        ("Platform",           "Cross-platform (Windows primary, Linux/macOS compatible)"),
        ("Interface",          "Graphical User Interface (CustomTkinter)"),
        ("Status",             "Active Development"),
        ("Repository",         "GitHub (Private)"),
        ("Latest Commit",      "April 16, 2026"),
        ("Author / Developer", "Maryam Inam"),
        ("Contact",            "ixxah.kaxmi@gmail.com"),
        ("Report Date",        "April 16, 2026"),
    ])

    add_heading(doc, "Problem Statement", 2)
    add_body(doc,
        "Phishing attacks remain the single most prevalent vector for data breaches and "
        "cyberattacks worldwide, accounting for over 36% of all breaches annually according to "
        "Verizon's Data Breach Investigations Report. Manual analysis of email headers—the "
        "primary forensic artifact for attributing and understanding phishing emails—is a "
        "time-consuming, expert-level task that requires knowledge of:")
    bullets = [
        "RFC 5322 email header standards and MIME encoding",
        "Email authentication protocols: SPF, DKIM, and DMARC",
        "IP geolocation and WHOIS domain registration interpretation",
        "Regex-based detection of phishing-specific language patterns",
        "Threat intelligence platforms and IOC extraction methodologies",
    ]
    for b in bullets:
        add_bullet(doc, b)

    add_body(doc,
        "This tool addresses that gap by providing an automated, GUI-driven platform that "
        "makes email forensic analysis accessible to security teams of all skill levels.")

    doc.add_page_break()

    # ================================================================
    # 3. OBJECTIVES & SCOPE
    # ================================================================
    add_heading(doc, "3. Objectives & Scope", 1)
    add_heading(doc, "Primary Objectives", 2)
    objectives = [
        ("Automate email header parsing", " and present data in human-readable form."),
        ("Trace email routing", " hop-by-hop from originating server to recipient."),
        ("Validate authentication", " — SPF, DKIM, DMARC — both from headers and live DNS."),
        ("Detect anomalies", " in timestamps, domain identities, and header consistency."),
        ("Analyse embedded URLs", " for display/href mismatches and domain reputation."),
        ("Scan attachments", " for dangerous file types and compute forensic hashes."),
        ("Integrate OSINT APIs", " for geolocation, IP abuse reputation, and malware detection."),
        ("Calculate risk score", " using a transparent, weighted scoring model."),
        ("Generate exportable reports", " in multiple formats for incident documentation."),
        ("Provide a professional GUI", " enabling use by non-expert security personnel."),
    ]
    for bold, normal in objectives:
        add_numbered(doc, normal, bold_prefix=bold)

    add_heading(doc, "Scope", 2)
    add_body(doc, "In Scope:", bold=True)
    in_scope = [
        "Analysis of .eml (RFC 5322) formatted email files",
        "Raw email header paste input",
        "Batch processing of multiple .eml files",
        "IPv4 geolocation and IP abuse reputation",
        "Domain WHOIS lookups and age analysis",
        "DNS record validation (SPF, DKIM, DMARC)",
        "Attachment hash computation and VirusTotal lookups",
        "HTML report and IOC export generation",
    ]
    for s in in_scope:
        add_bullet(doc, s)

    add_body(doc, "Out of Scope:", bold=True)
    out_scope = [
        "Live email inbox monitoring or IMAP/POP3 integration",
        "IPv6 geolocation",
        "Automated email response or quarantine actions",
        "Machine-learning-based classification (rule-based only)",
        "Mobile platform support",
    ]
    for s in out_scope:
        add_bullet(doc, s)

    doc.add_page_break()

    # ================================================================
    # 4. SYSTEM ARCHITECTURE
    # ================================================================
    add_heading(doc, "4. System Architecture", 1)
    add_body(doc,
        "The application follows a clean two-layer separation of concerns: a headless "
        "backend engine and a presentation GUI layer. This design enables the backend "
        "to be used independently via command line or imported as a Python library.")

    add_heading(doc, "Architectural Diagram", 2)
    add_code_block(doc, """\
┌─────────────────────────────────────────────────────────────────┐
│                   USER INTERFACE LAYER                          │
│              forensic_gui.py  (CustomTkinter)                   │
│  ┌──────────┐ ┌───────────┐ ┌──────────┐ ┌───────────────────┐ │
│  │ Sidebar  │ │ Tab Bar   │ │ Progress │ │  Export / Reports │ │
│  │ Controls │ │ (10 tabs) │ │   Bar    │ │                   │ │
│  └──────────┘ └───────────┘ └──────────┘ └───────────────────┘ │
└────────────────────────┬────────────────────────────────────────┘
                         │  calls / callbacks (threading)
┌────────────────────────▼────────────────────────────────────────┐
│                 BUSINESS LOGIC LAYER                            │
│          email_forensic_analyzer.py (EmailForensicAnalyzer)     │
│  ┌──────────┐ ┌───────────┐ ┌──────────┐ ┌───────────────────┐ │
│  │  Parser  │ │  Routing  │ │  Auth    │ │  Risk Scoring     │ │
│  │ (email)  │ │  (hops)   │ │ SPF/DKIM │ │  IOC Extraction   │ │
│  └──────────┘ └───────────┘ └──────────┘ └───────────────────┘ │
└────────────────────────┬────────────────────────────────────────┘
                         │  HTTP / DNS queries
┌────────────────────────▼────────────────────────────────────────┐
│                  INTEGRATION LAYER                              │
│  ┌──────────────┐ ┌───────────────┐ ┌──────────────────────┐   │
│  │ ip-api.com   │ │  AbuseIPDB    │ │     VirusTotal        │   │
│  │ (Geo/ISP)    │ │  (IP Abuse)   │ │  (File Hash Scan)     │   │
│  └──────────────┘ └───────────────┘ └──────────────────────┘   │
│  ┌──────────────┐ ┌───────────────┐                             │
│  │ DNS Resolver │ │  WHOIS        │                             │
│  │ (dnspython)  │ │ (python-whois)│                             │
│  └──────────────┘ └───────────────┘                             │
└─────────────────────────────────────────────────────────────────┘""")

    add_heading(doc, "Threading Model", 2)
    add_body(doc,
        "To prevent the GUI from freezing during network-dependent operations, the tool "
        "employs a multi-threaded architecture:")
    add_section_table(doc,
        ["Thread", "Responsibility", "Triggers UI Update Via"],
        [
            ["Main Thread",        "GUI rendering, user interaction, file I/O", "Direct"],
            ["Geolocation Thread", "ip-api.com queries for all routing IPs",     "self.after(0, callback)"],
            ["WHOIS Thread",       "Domain registration and age lookup",          "self.after(0, callback)"],
            ["ThreatIntel Thread", "DNS validation, phishing patterns, AbuseIPDB","self.after(0, callback)"],
            ["Batch Thread",       "Multi-file .eml processing loop",             "self.after(0, callback)"],
            ["VirusTotal Thread",  "Hash-based malware lookup (optional)",        "self.after(0, callback)"],
        ])

    add_heading(doc, "Data Caching Strategy", 2)
    add_body(doc,
        "All analysis results are stored in a result cache dictionary keyed by file path. "
        "This enables the user to toggle between light and dark themes without re-running "
        "expensive API queries. The cache is maintained in memory for the application session "
        "and cleared on application exit.")

    doc.add_page_break()

    # ================================================================
    # 5. TECHNOLOGY STACK
    # ================================================================
    add_heading(doc, "5. Technology Stack", 1)

    add_heading(doc, "5.1 Core Language", 2)
    add_info_table(doc, [
        ("Language",   "Python 3.10+"),
        ("Features Used", "Type hints (PEP 484), f-strings, dataclasses, threading, contextlib"),
        ("Code Style", "PEP 8 compliant with full docstring coverage"),
    ])

    add_heading(doc, "5.2 Third-Party Dependencies", 2)
    add_section_table(doc,
        ["Package", "Version", "Purpose"],
        [
            ["customtkinter",  "≥ 5.2.0", "Modern, themeable GUI framework (replaces standard tkinter)"],
            ["beautifulsoup4", "≥ 4.12.0","HTML/XML parsing for URL extraction from email bodies"],
            ["dnspython",      "≥ 2.4.0", "DNS record lookups: SPF TXT, DKIM selector, DMARC TXT"],
            ["python-whois",   "≥ 0.9.0", "Domain WHOIS queries for registration date and registrar"],
            ["requests",       "≥ 2.28.0","HTTP client for REST API calls to ip-api, AbuseIPDB, VirusTotal"],
        ])

    add_heading(doc, "5.3 Python Standard Library Modules", 2)
    add_section_table(doc,
        ["Module", "Usage"],
        [
            ["email, email.policy, email.utils", "RFC 5322 email parsing, header access, address decoding"],
            ["re",                               "Regex-based pattern matching (50+ compiled patterns)"],
            ["ipaddress",                        "IP address validation, classification (private/public/reserved)"],
            ["hashlib",                          "MD5 and SHA-256 digest computation for attachment forensics"],
            ["json, csv",                        "IOC serialisation for export"],
            ["threading",                        "Background task execution to prevent GUI blocking"],
            ["tkinter",                          "GUI foundation, file dialogs (via customtkinter wrapper)"],
            ["pathlib",                          "Cross-platform file path handling"],
            ["datetime, timezone",               "Timestamp parsing, UTC conversion, domain age calculation"],
            ["urllib.parse",                     "URL decoding and domain extraction"],
            ["html.parser",                      "Fallback HTML parsing"],
        ])

    add_heading(doc, "5.4 External APIs", 2)
    add_section_table(doc,
        ["API", "Provider", "Auth Required", "Purpose"],
        [
            ["ip-api.com",  "ip-api",   "No (free tier)",  "IP geolocation, ISP, ASN, country, city"],
            ["AbuseIPDB",   "AbuseIPDB","Yes (API key)",    "IP abuse confidence score, report count"],
            ["VirusTotal",  "VirusTotal","Yes (API key)",   "File hash reputation and malware detection"],
        ])

    doc.add_page_break()

    # ================================================================
    # 6. CORE FEATURES & FUNCTIONALITY
    # ================================================================
    add_heading(doc, "6. Core Features & Functionality", 1)

    feature_sections = [
        ("6.1 Email File Loading", [
            "File browser dialog filtered to .eml extension",
            "Raw email / header paste via modal popup dialog",
            "Automatic MIME structure detection (simple vs. multipart)",
            "Recent analyses list maintained in the sidebar (last 5 analyses)",
        ]),
        ("6.2 Header Extraction & Metadata", [
            "Extracts: From, To, CC, Subject, Date, Message-ID, Return-Path, Reply-To",
            "Handles RFC 2047 encoded-word decoding for non-ASCII header values",
            "Extracts 11 forensically significant X-Headers",
            "Displays raw header value alongside decoded interpretation",
        ]),
        ("6.3 Email Routing Path Analysis", [
            "Parses all Received: headers (chronological reversal from RFC 5322 order)",
            "Extracts from/by domain clauses and originating IP per hop",
            "Classifies IPs as public, private (RFC 1918), loopback, or reserved",
            "Computes inter-hop time deltas and flags anomalies",
            "Detects: time-travel (negative delta), excessive delays (>4 hours), future-dated timestamps",
        ]),
        ("6.4 Authentication Verification", [
            "Parses Authentication-Results header for SPF, DKIM, and DMARC results",
            "Falls back to Received-SPF header when Authentication-Results is absent",
            "Normalises result strings (pass/fail/softfail/neutral/none/permerror/temperror)",
            "Flags emails as suspicious if SPF or DKIM result is fail or softfail",
            "Performs live DNS queries to validate SPF TXT, DKIM selector, and DMARC TXT records",
            "Tries 8 common DKIM selectors: default, google, selector1, selector2, s1, s2, k1, mail",
        ]),
        ("6.5 Header Anomaly Detection", [
            "From vs. Reply-To domain mismatch",
            "From vs. Return-Path domain mismatch",
            "From vs. DKIM signing domain (d= parameter) mismatch",
            "Timestamp sequence validation across all hops",
            "Date header validation against current UTC time",
        ]),
        ("6.6 URL & Link Analysis", [
            "Parses HTML body with BeautifulSoup to extract all <a href=...> anchor tags",
            "Compares display text domain vs. actual href domain for mismatch detection",
            "Falls back to plain-text regex URL extraction for non-HTML emails",
            "URL deduplication using a seen-set to prevent duplicate entries",
            "Performs WHOIS domain age lookup on extracted domains",
        ]),
        ("6.7 Attachment Analysis", [
            "Walks MIME tree to detect all attachments regardless of nesting depth",
            "Extracts: filename, MIME type, size in bytes",
            "Computes MD5 and SHA-256 cryptographic hashes for each attachment",
            "Flags 48 risky file extensions (executables, scripts, macros, archives, disk images)",
            "Optional VirusTotal integration for hash-based malware scanning",
        ]),
        ("6.8 Risk Scoring", [
            "Produces a 0–100 integer risk score from 8 weighted signals",
            "Maps score to threat levels: Critical (75–100), High (50–74), Medium (25–49), Low (0–24)",
            "Provides a breakdown dictionary explaining contribution of each signal",
            "Displays score as an animated semicircular gauge on the Dashboard tab",
        ]),
        ("6.9 IOC Export", [
            "Aggregates all indicators: IP addresses, domains, URLs, email addresses, file hashes",
            "Exports to JSON (structured) or CSV (flat) format",
            "Clipboard copy support for individual indicators",
            "IOC tab displays categorised lists with copy buttons",
        ]),
        ("6.10 Report Generation", [
            "Self-contained HTML reports with embedded CSS (no external dependencies)",
            "Batch analysis: process entire folder, generate summary comparison table",
            "Professional printable layout with colour-coded risk indicators",
            "Timestamps and analyst notes supported in report metadata",
        ]),
    ]
    for title, bullets in feature_sections:
        add_heading(doc, title, 2)
        for b in bullets:
            add_bullet(doc, b)

    doc.add_page_break()

    # ================================================================
    # 7. BACKEND ANALYSIS ENGINE
    # ================================================================
    add_heading(doc, "7. Backend Analysis Engine", 1)
    add_body(doc,
        "The backend is encapsulated in the EmailForensicAnalyzer class "
        "(email_forensic_analyzer.py, 1,627 lines). It is deliberately designed as a "
        "headless module with no GUI dependencies, enabling command-line use and unit testing.")

    add_heading(doc, "7.1 Class Interface", 2)
    add_code_block(doc, """\
class EmailForensicAnalyzer:
    def __init__(self, email_content: str)
    def extract_basic_metadata(self)      -> dict
    def extract_routing_path(self)         -> list[dict]
    def check_authentication(self)         -> dict
    def check_header_anomalies(self)       -> dict
    def extract_urls(self)                 -> list[dict]
    def extract_attachments(self)          -> list[dict]
    def get_ip_geolocation(ip)             -> dict          # staticmethod
    def get_domain_reputation(domain)      -> dict          # staticmethod
    def check_ip_abuse(ip, api_key)        -> dict          # staticmethod
    def check_virustotal(hash, api_key)    -> dict          # staticmethod
    def detect_phishing_patterns(self)     -> dict
    def calculate_risk_score(self, ...)    -> dict
    def extract_iocs(self, ...)            -> dict""")

    add_heading(doc, "7.2 Key Method Descriptions", 2)
    add_section_table(doc,
        ["Method", "Returns", "Description"],
        [
            ["extract_basic_metadata()",  "dict",       "Extracts From, To, Subject, Date, Message-ID, Return-Path"],
            ["extract_routing_path()",    "list[dict]", "Parses Received headers into ordered hop list"],
            ["check_authentication()",    "dict",       "SPF/DKIM/DMARC results + DNS record validation"],
            ["check_header_anomalies()",  "dict",       "Domain mismatches, timestamp issues, X-header values"],
            ["extract_urls()",            "list[dict]", "All URLs with display/href mismatch flags"],
            ["extract_attachments()",     "list[dict]", "Attachment metadata, hashes, risky extension flag"],
            ["detect_phishing_patterns()","dict",       "Matched urgency, credential, and impersonation patterns"],
            ["calculate_risk_score()",    "dict",       "Weighted 0-100 score with breakdown"],
            ["extract_iocs()",            "dict",       "Aggregated IOCs by category"],
            ["get_ip_geolocation()",      "dict",       "Static: calls ip-api.com, returns geo/ISP/ASN"],
            ["get_domain_reputation()",   "dict",       "Static: WHOIS lookup, domain age, registrar"],
            ["check_ip_abuse()",          "dict",       "Static: AbuseIPDB confidence score"],
            ["check_virustotal()",        "dict",       "Static: VirusTotal file hash reputation"],
        ])

    add_heading(doc, "7.3 Return Data Structures", 2)
    add_heading(doc, "extract_routing_path() — List Item", 3)
    add_code_block(doc, """\
{
  "hop":       int,    # chronological hop number (1 = originating server)
  "from":      str,    # sending server hostname/domain
  "by":        str,    # receiving server hostname/domain
  "ip":        str,    # originating IP address (may be None for private hops)
  "timestamp": str     # ISO-format UTC timestamp string
}""")

    add_heading(doc, "calculate_risk_score() — Return Structure", 3)
    add_code_block(doc, """\
{
  "score":     int,       # 0–100
  "level":     str,       # "Critical" | "High" | "Medium" | "Low"
  "breakdown": {
    "auth_failure":    (int_points, str_reason),
    "url_mismatch":    (int_points, str_reason),
    "risky_attachment":(int_points, str_reason),
    "young_domain":    (int_points, str_reason),
    "abusive_ip":      (int_points, str_reason),
    "urgency_language":(int_points, str_reason),
    "credential_lang": (int_points, str_reason),
    "impersonation":   (int_points, str_reason),
  }
}""")

    doc.add_page_break()

    # ================================================================
    # 8. GRAPHICAL USER INTERFACE
    # ================================================================
    add_heading(doc, "8. Graphical User Interface (GUI)", 1)
    add_body(doc,
        "The GUI (forensic_gui.py, 2,214 lines) is built with CustomTkinter, a modern "
        "Python GUI library that extends standard tkinter with rounded widgets, high-DPI "
        "support, and built-in dark/light mode theming. The interface is designed around "
        "a sidebar-plus-tabs layout pattern common in professional desktop applications.")

    add_heading(doc, "8.1 Window Layout", 2)
    add_code_block(doc, """\
┌──────────────────────────────────────────────────────────────────┐
│  TITLE BAR                                                       │
├──────────────────┬───────────────────────────────────────────────┤
│  SIDEBAR (220px) │  MAIN CONTENT AREA (flexible)                 │
│                  │  ┌──────────────────────────────────────────┐ │
│  [Select File]   │  │  Risk Warning Banner (conditional)       │ │
│  [Paste Headers] │  ├──────────────────────────────────────────┤ │
│  [Batch Analyze] │  │  Tab Navigation Bar (10 tabs)            │ │
│  ─────────────   │  ├──────────────────────────────────────────┤ │
│  [API Keys ▼]    │  │                                          │ │
│   • AbuseIPDB    │  │  Active Tab Content (scrollable)         │ │
│   • VirusTotal   │  │                                          │ │
│  ─────────────   │  │                                          │ │
│  Recent (5)      │  └──────────────────────────────────────────┘ │
│  [Analysis 1]    │  ┌──────────────────────────────────────────┐ │
│  [Analysis 2]    │  │  Progress Bar (8-phase, hidden at rest)  │ │
│  ─────────────   │  ├──────────────────────────────────────────┤ │
│  Summary Card    │  │  Status Bar (file path + current action) │ │
│  ─────────────   │  └──────────────────────────────────────────┘ │
│  [Theme Toggle]  │                                               │
│  [Export Report] │                                               │
└──────────────────┴───────────────────────────────────────────────┘""")

    add_heading(doc, "8.2 Tab Structure and Content", 2)
    add_section_table(doc,
        ["Tab #", "Tab Name", "Key Content"],
        [
            ["1",  "Dashboard",        "Risk gauge (0-100), verdict badge, quick stats, key findings, sender identity card"],
            ["2",  "Metadata",         "All standard headers: From, To, Subject, Date, Message-ID, Reply-To, Return-Path"],
            ["3",  "Header Analysis",  "Domain identity mismatches, timestamp anomalies, forensic X-header values"],
            ["4",  "Routing Path",     "Chronological hop table: hop #, from, by, IP, timestamp, inter-hop delta"],
            ["5",  "Authentication",   "SPF / DKIM / DMARC results with PASS/FAIL/WARN badges + DNS record validation"],
            ["6",  "Geolocation",      "Originating IP map data: country, city, ISP, ASN; all hop IPs"],
            ["7",  "URLs & Links",     "All extracted URLs; mismatch alerts; domain age for each domain"],
            ["8",  "Attachments",      "File metadata, MD5/SHA-256 hashes, risk extension flag, VirusTotal result"],
            ["9",  "Threat Intel",     "Risk score breakdown bar chart, DNS records, phishing pattern matches, AbuseIPDB"],
            ["10", "IOC Export",       "Categorised IOC list (IPs, domains, URLs, emails, hashes) with export buttons"],
        ])

    add_heading(doc, "8.3 Visual Design System", 2)
    add_section_table(doc,
        ["Element", "Light Mode", "Dark Mode"],
        [
            ["Background",           "#F4F7FA (light grey)",  "#1E2B3C (dark navy)"],
            ["Primary Accent",       "#2980B9 (blue)",        "#3498DB (bright blue)"],
            ["Danger / Fail",        "#C0392B (red)",         "#E74C3C (bright red)"],
            ["Safe / Pass",          "#1A8A4A (green)",       "#2ECC71 (bright green)"],
            ["Warning",              "#E67E22 (orange)",      "#F39C12 (bright orange)"],
            ["Headings",             "#0D2B45 (navy)",        "#ECF0F1 (near-white)"],
            ["Card Borders",         "#D0D7DE",               "#2C3E50"],
            ["Typography (body)",    "Calibri / Segoe UI 11pt","Same"],
        ])

    add_heading(doc, "8.4 Risk Gauge Widget", 2)
    add_body(doc,
        "The Dashboard tab features a custom-drawn semicircular canvas gauge widget. "
        "It renders a coloured arc (0–180 degrees) mapping the risk score (0–100) to a "
        "colour gradient from green through yellow to red. The needle animates from 0 to "
        "the final score when analysis completes. The gauge is re-drawn on theme switch "
        "without re-running analysis.")

    add_heading(doc, "8.5 Progress Bar Phases", 2)
    add_section_table(doc,
        ["Phase #", "Label", "Operations"],
        [
            ["1", "Parsing Email",         "Load .eml file, detect MIME structure"],
            ["2", "Extracting Headers",    "Metadata, X-headers, routing path"],
            ["3", "Checking Authentication","SPF/DKIM/DMARC header parsing"],
            ["4", "Analysing Headers",     "Anomaly detection, domain comparisons"],
            ["5", "Scanning URLs",         "HTML parsing, URL extraction, mismatch detection"],
            ["6", "Analysing Attachments", "MIME walk, hash computation"],
            ["7", "Threat Intelligence",   "Async: DNS, WHOIS, AbuseIPDB, geolocation"],
            ["8", "Calculating Risk Score","Weighted scoring, IOC aggregation"],
        ])

    doc.add_page_break()

    # ================================================================
    # 9. THREAT INTELLIGENCE INTEGRATION
    # ================================================================
    add_heading(doc, "9. Threat Intelligence Integration", 1)

    add_heading(doc, "9.1 IP Geolocation — ip-api.com", 2)
    add_info_table(doc, [
        ("Endpoint",          "http://ip-api.com/json/{ip}?fields=status,country,city,isp,org,as,query"),
        ("Authentication",    "None required (free tier: 45 requests/minute)"),
        ("Private IP Handling","Private/reserved IPs are detected locally and skipped"),
        ("Data Returned",     "Country, city, ISP name, ASN, query IP"),
        ("Timeout",           "5 seconds"),
    ])

    add_heading(doc, "9.2 IP Abuse Reputation — AbuseIPDB", 2)
    add_info_table(doc, [
        ("Endpoint",         "https://api.abuseipdb.com/api/v2/check"),
        ("Authentication",   "API key required (X-Auth-Token header)"),
        ("Data Returned",    "Abuse confidence score (0-100%), total reports, ISP, usage type"),
        ("Risk Threshold",   "Score ≥ 25% triggers risk signal (+10 points to risk score)"),
        ("Timeout",          "5 seconds"),
    ])

    add_heading(doc, "9.3 Malware Detection — VirusTotal", 2)
    add_info_table(doc, [
        ("Endpoint",         "https://www.virustotal.com/api/v3/files/{sha256_hash}"),
        ("Authentication",   "API key required (x-apikey header)"),
        ("Data Returned",    "Malicious verdict count, total engine count, detection names"),
        ("Trigger",          "SHA-256 hash of each attachment"),
        ("Timeout",          "5 seconds"),
    ])

    add_heading(doc, "9.4 DNS Record Validation — dnspython", 2)
    add_body(doc,
        "The tool performs three live DNS lookups per domain to validate the presence "
        "and content of email authentication records:")
    add_section_table(doc,
        ["Record Type", "Query Target", "Validation"],
        [
            ["SPF",   "{domain}",                  "TXT record starting with v=spf1"],
            ["DKIM",  "{selector}._domainkey.{domain}", "TXT record containing p= (public key)"],
            ["DMARC", "_dmarc.{domain}",            "TXT record starting with v=DMARC1"],
        ])
    add_body(doc,
        "DKIM selector discovery tries: default, google, selector1, selector2, s1, s2, k1, mail "
        "before reporting no DKIM record found.")

    add_heading(doc, "9.5 Domain WHOIS — python-whois", 2)
    add_info_table(doc, [
        ("Library",          "python-whois ≥ 0.9.0"),
        ("Data Extracted",   "creation_date, registrar, expiration_date"),
        ("Domain Age Logic", "Computed as (today - creation_date) in days"),
        ("Young Domain Flag","< 30 days since registration triggers risk signal (+15 points)"),
    ])

    doc.add_page_break()

    # ================================================================
    # 10. RISK SCORING METHODOLOGY
    # ================================================================
    add_heading(doc, "10. Risk Scoring Methodology", 1)
    add_body(doc,
        "The risk scoring engine produces a deterministic, transparent score by summing "
        "the contributions of eight independently evaluated signals. The maximum achievable "
        "score is 100 points.")

    add_heading(doc, "10.1 Scoring Signals", 2)
    add_section_table(doc,
        ["Signal", "Points", "Condition", "Rationale"],
        [
            ["Authentication Failure", "25", "SPF or DKIM result is fail/softfail",    "Strongest indicator of spoofed sender"],
            ["URL Display Mismatch",   "20", "Any URL has display≠href domain",         "Classic phishing deception technique"],
            ["Risky Attachment",       "15", "Any attachment has dangerous extension",  "Common malware delivery method"],
            ["Young Domain",           "15", "Sending domain < 30 days old",            "Phishing domains are registered fresh"],
            ["Abusive IP",             "10", "IP abuse confidence ≥ 25%",               "Previously reported malicious IP"],
            ["Urgency Language",       "5",  "Urgency pattern matched in body/subject", "Psychological manipulation tactic"],
            ["Credential Language",    "5",  "Credential-harvesting terms matched",     "Explicit credential theft attempt"],
            ["Brand Impersonation",    "5",  "Brand name impersonation matched",        "Social engineering via trusted brands"],
        ])

    add_heading(doc, "10.2 Threat Level Mapping", 2)
    add_section_table(doc,
        ["Score Range", "Threat Level", "Recommended Action"],
        [
            ["75–100", "CRITICAL", "Immediate quarantine; full incident response; block sender domain/IP"],
            ["50–74",  "HIGH",     "Do not interact; escalate to security team; preserve email as evidence"],
            ["25–49",  "MEDIUM",   "Caution advised; verify sender through out-of-band channel"],
            ["0–24",   "LOW",      "Likely benign; no immediate action required"],
        ])

    add_callout(doc,
        "Design Note: Signals are additive and non-exclusive. An email that fails "
        "authentication AND contains a risky attachment AND uses urgency language would "
        "score 25+15+5=45 (Medium), while one that also has a URL mismatch and a young "
        "domain would score 25+20+15+15=75 (Critical).",
        "info")

    doc.add_page_break()

    # ================================================================
    # 11. DATA FLOW & PROCESSING PIPELINE
    # ================================================================
    add_heading(doc, "11. Data Flow & Processing Pipeline", 1)

    add_heading(doc, "11.1 Single Email Analysis Pipeline", 2)
    add_code_block(doc, """\
  User Input
      │
      ▼
  [File Dialog / Paste Dialog]
      │
      ▼
  email.message_from_string()   ← RFC 5322 parser
      │
      ├─► extract_basic_metadata()       [synchronous]
      ├─► extract_routing_path()         [synchronous]
      ├─► check_authentication()         [synchronous + DNS queries]
      ├─► check_header_anomalies()       [synchronous]
      ├─► extract_urls()                 [synchronous + WHOIS async]
      ├─► extract_attachments()          [synchronous + VT async]
      │
      ▼                                   [background threads]
  ┌─────────────────────────────────────────────────┐
  │  Thread 1: get_ip_geolocation()  for each hop   │
  │  Thread 2: get_domain_reputation() for domains  │
  │  Thread 3: detect_phishing_patterns()           │
  │            check_ip_abuse()                     │
  └─────────────────────────────────────────────────┘
      │  (callbacks via self.after(0, fn))
      ▼
  calculate_risk_score()
      │
      ▼
  extract_iocs()
      │
      ├─► Populate GUI Tabs
      ├─► Update Sidebar Summary
      └─► Enable Export Buttons""")

    add_heading(doc, "11.2 Batch Analysis Pipeline", 2)
    add_code_block(doc, """\
  User selects folder
      │
      ▼
  glob("*.eml") → file list
      │
      ▼
  Background Thread:
    for each .eml file:
      EmailForensicAnalyzer(content).run_all()
      results.append(summary_dict)
      self.after(0, update_progress)
      │
      ▼
  Generate HTML Summary Table
      │
      ▼
  Save to {folder}/batch_report_{timestamp}.html""")

    doc.add_page_break()

    # ================================================================
    # 12. ALGORITHMS & TECHNICAL IMPLEMENTATION
    # ================================================================
    add_heading(doc, "12. Algorithms & Technical Implementation", 1)

    add_heading(doc, "12.1 Routing Path Extraction Algorithm", 2)
    add_code_block(doc, """\
def extract_routing_path(self) -> list[dict]:
    received_headers = self.msg.get_all("Received", [])
    # RFC 5322: headers stored newest-first → reverse for chronological order
    received_headers.reverse()

    hops = []
    for i, header in enumerate(received_headers, start=1):
        # Collapse folded whitespace
        header = re.sub(r'\\s+', ' ', header)

        from_match = re.search(r'from\\s+(\\S+)', header)
        by_match   = re.search(r'by\\s+(\\S+)',   header)
        ip_match   = re.search(r'\\[(\\d+\\.\\d+\\.\\d+\\.\\d+)\\]', header)

        # Extract timestamp after semicolon
        ts_match   = re.search(r';\\s*(.+)$', header)
        timestamp  = parse_email_date(ts_match.group(1)) if ts_match else None

        ip = ip_match.group(1) if ip_match else None
        if ip and ipaddress.ip_address(ip).is_private:
            ip = None   # skip private IPs for geolocation

        hops.append({
            "hop": i,
            "from": from_match.group(1) if from_match else "unknown",
            "by":   by_match.group(1)   if by_match   else "unknown",
            "ip":   ip,
            "timestamp": timestamp.isoformat() if timestamp else "unparseable"
        })
    return hops""")

    add_heading(doc, "12.2 Authentication Verification Algorithm", 2)
    add_code_block(doc, """\
def check_authentication(self) -> dict:
    auth_results = self.msg.get("Authentication-Results", "")
    result = {"spf": "none", "dkim": "none", "dmarc": "none", "is_suspicious": False}

    if auth_results:
        spf_match  = re.search(r'spf=(\\w+)',  auth_results, re.I)
        dkim_match = re.search(r'dkim=(\\w+)', auth_results, re.I)
        dmarc_match= re.search(r'dmarc=(\\w+)',auth_results, re.I)
        result["spf"]   = spf_match.group(1).lower()  if spf_match  else "none"
        result["dkim"]  = dkim_match.group(1).lower() if dkim_match else "none"
        result["dmarc"] = dmarc_match.group(1).lower()if dmarc_match else "none"
    else:
        # Fallback to Received-SPF header
        received_spf = self.msg.get("Received-SPF", "")
        if received_spf:
            match = re.match(r'^(\\w+)', received_spf.strip())
            result["spf"] = match.group(1).lower() if match else "none"

    suspicious = {"fail", "softfail"}
    result["is_suspicious"] = bool(
        {result["spf"], result["dkim"]} & suspicious
    )
    return result""")

    add_heading(doc, "12.3 Phishing Pattern Matching", 2)
    add_body(doc,
        "All patterns are precompiled at module load time for performance. "
        "The search is case-insensitive and operates on the concatenated subject + body text.")
    add_code_block(doc, """\
URGENCY_PATTERNS = [re.compile(p, re.I) for p in [
    r'\\burgent\\b', r'\\bact now\\b', r'\\bimmediate action\\b',
    r'\\baccount.*suspend', r'\\bverify your\\b', r'\\bexpir',
    r'\\blimited time\\b', r'\\bclick immediately\\b', ...
]]

def detect_phishing_patterns(self) -> dict:
    text = (self.get_subject() + " " + self.get_body_text()).lower()
    return {
        "urgency":       list({m.group() for p in URGENCY_PATTERNS
                                for m in [p.search(text)] if m}),
        "credential":    list({m.group() for p in CREDENTIAL_PATTERNS ...}),
        "impersonation": list({m.group() for p in BRAND_PATTERNS ...}),
    }""")

    add_heading(doc, "12.4 URL Display/Href Mismatch Detection", 2)
    add_code_block(doc, """\
def extract_urls(self) -> list[dict]:
    urls, seen = [], set()
    soup = BeautifulSoup(html_body, "html.parser")
    for a in soup.find_all("a", href=True):
        href = a["href"].strip()
        display = a.get_text(strip=True)
        if href in seen: continue
        seen.add(href)
        href_domain    = urlparse(href).netloc.lower()
        display_domain = urlparse(display).netloc.lower() if "http" in display else ""
        mismatch = bool(
            display_domain and
            href_domain and
            display_domain != href_domain
        )
        urls.append({"url": href, "display_text": display,
                     "domain": href_domain, "mismatch": mismatch})
    return urls""")

    doc.add_page_break()

    # ================================================================
    # 13. SECURITY CONSIDERATIONS
    # ================================================================
    add_heading(doc, "13. Security Considerations", 1)

    add_heading(doc, "13.1 Input Handling", 2)
    bullets = [
        "Email content is never executed — only parsed as text/MIME",
        "BeautifulSoup parses HTML in a sandboxed DOM context with no JavaScript execution",
        "File paths are validated via pathlib before file operations",
        "API keys are stored in memory only (not written to disk or logs)",
        "WHOIS and DNS queries use timeouts to prevent hanging on malicious input",
    ]
    for b in bullets:
        add_bullet(doc, b)

    add_heading(doc, "13.2 API Security", 2)
    bullets = [
        "AbuseIPDB and VirusTotal API keys are passed as HTTP headers, not query parameters",
        "All external API calls use a 5-second timeout to mitigate slow-loris attacks",
        "The application does not store or cache API keys between sessions",
        "Private IP addresses are never sent to external geolocation APIs",
    ]
    for b in bullets:
        add_bullet(doc, b)

    add_heading(doc, "13.3 Report Generation Security", 2)
    bullets = [
        "HTML reports use html.escape() for all user-derived values to prevent stored XSS",
        "Reports are self-contained (no external script/CSS sources loaded) — safe for offline use",
        "Attachment hashes are computed locally — no file content is uploaded externally",
    ]
    for b in bullets:
        add_bullet(doc, b)

    add_callout(doc,
        "Privacy Note: The tool sends IP addresses to ip-api.com (geolocation) and "
        "optionally to AbuseIPDB. Users operating in privacy-sensitive environments should "
        "review the privacy policies of these services before enabling API lookups.",
        "warning")

    doc.add_page_break()

    # ================================================================
    # 14. TESTING & SAMPLE DATA
    # ================================================================
    add_heading(doc, "14. Testing & Sample Data", 1)
    add_heading(doc, "14.1 Included Sample Files", 2)
    add_section_table(doc,
        ["File", "Size", "Purpose"],
        [
            ["email.eml",  "534 lines", "Full production-style phishing email with all header types, HTML body, and attachments"],
            ["sample.eml", "61 lines",  "Minimal email for quick UI and parsing smoke-tests"],
        ])

    add_heading(doc, "14.2 Test Coverage Areas", 2)
    test_areas = [
        "SPF fail with valid Received-SPF fallback",
        "DKIM pass with selector discovery",
        "Multi-hop routing path with private IP filtering",
        "URL display/href mismatch in HTML body",
        "Dangerous attachment extension detection (.exe, .docm)",
        "Phishing pattern matching (urgency + credential + brand)",
        "Domain mismatch: From vs. Reply-To",
        "Young domain WHOIS flag",
        "Batch analysis with mixed-risk emails",
        "Theme switching with cached results",
    ]
    for t in test_areas:
        add_bullet(doc, t)

    doc.add_page_break()

    # ================================================================
    # 15. INSTALLATION & DEPLOYMENT
    # ================================================================
    add_heading(doc, "15. Installation & Deployment", 1)

    add_heading(doc, "15.1 System Requirements", 2)
    add_section_table(doc,
        ["Requirement", "Minimum", "Recommended"],
        [
            ["Python",       "3.10",           "3.12+"],
            ["RAM",          "256 MB",          "512 MB+"],
            ["Disk Space",   "50 MB",           "100 MB"],
            ["Network",      "Optional",        "Required for API features"],
            ["OS",           "Windows 10",      "Windows 11 / Ubuntu 22.04 / macOS 12+"],
            ["Display",      "1280×720",        "1920×1080"],
        ])

    add_heading(doc, "15.2 Installation Steps", 2)
    steps = [
        ("Clone or download", " the repository to your local machine."),
        ("Create a virtual environment:", " python -m venv venv"),
        ("Activate the environment:", " Windows: venv\\Scripts\\activate | Linux/Mac: source venv/bin/activate"),
        ("Install dependencies:", " pip install -r requirements.txt"),
        ("Launch the application:", " python forensic_gui.py"),
    ]
    for bold, normal in steps:
        add_numbered(doc, normal, bold_prefix=bold)

    add_heading(doc, "15.3 Optional API Key Configuration", 2)
    add_body(doc,
        "API keys are entered directly in the GUI's expandable \"API Keys\" panel in the sidebar. "
        "No configuration file or environment variable setup is required. Keys are retained "
        "in memory for the current session only.")

    doc.add_page_break()

    # ================================================================
    # 16. KNOWN LIMITATIONS & FUTURE ENHANCEMENTS
    # ================================================================
    add_heading(doc, "16. Known Limitations & Future Enhancements", 1)

    add_heading(doc, "16.1 Current Limitations", 2)
    limitations = [
        "IPv6 addresses are detected but not geolocated (ip-api.com limited support)",
        "DKIM cryptographic signature verification is not performed (result relies on mail server's Authentication-Results header)",
        "No IMAP/POP3 live mailbox integration — .eml files must be exported manually",
        "WHOIS data accuracy depends on registrar data completeness and python-whois parsing",
        "VirusTotal API has a 4 requests/minute rate limit on the free tier",
        "Batch analysis runs sequentially (no parallel multi-threading per file)",
        "No persistent database for historical analysis storage",
    ]
    for l in limitations:
        add_bullet(doc, l)

    add_heading(doc, "16.2 Planned Future Enhancements", 2)
    enhancements = [
        ("IMAP Integration:", " Connect directly to email inboxes for one-click analysis"),
        ("ML Classification:", " Train a lightweight model on IOC features for probabilistic scoring"),
        ("DKIM Crypto Verification:", " Verify DKIM signatures using the public key from DNS"),
        ("PostgreSQL Backend:", " Persistent storage of analyses for trend reporting"),
        ("STIX/TAXII Export:", " Export IOCs in STIX 2.1 format for SIEM/SOAR integration"),
        ("IPv6 Geolocation:", " Extend geolocation to IPv6 routing hops"),
        ("Parallel Batch Processing:", " Use ThreadPoolExecutor for concurrent multi-file analysis"),
        ("Web Interface:", " Flask/FastAPI REST API backend with React frontend"),
        ("Plugin Architecture:", " Allow custom detection rule plugins"),
        ("Automated Reporting Schedule:", " Schedule periodic batch analysis runs"),
    ]
    for bold, normal in enhancements:
        add_bullet(doc, normal, bold_prefix=bold)

    doc.add_page_break()

    # ================================================================
    # 17. CONCLUSION
    # ================================================================
    add_heading(doc, "17. Conclusion", 1)
    add_body(doc,
        "The Email Header Phishing Investigation Tool represents a comprehensive, "
        "professionally engineered solution to one of cybersecurity's most persistent "
        "challenges. By automating the forensic analysis of email headers, the tool "
        "dramatically reduces the expertise barrier and time investment required to "
        "investigate phishing attacks.")
    add_body(doc,
        "The application's clean two-layer architecture (backend engine + GUI frontend), "
        "robust threading model, and extensive threat intelligence integrations position it "
        "as a production-ready tool suitable for SOC (Security Operations Centre) teams, "
        "digital forensics practitioners, and security researchers.")
    add_body(doc,
        "The combination of deterministic rule-based scoring (transparent and auditable), "
        "real-time OSINT API integration, and professional export capabilities creates a "
        "tool that bridges the gap between raw email forensics and actionable threat "
        "intelligence reporting.")
    add_body(doc,
        "With a roadmap including machine-learning classification, IMAP integration, and "
        "STIX/TAXII export, the tool is positioned for evolution into a full-featured "
        "email security investigation platform.")

    add_callout(doc,
        "The project demonstrates strong software engineering fundamentals: separation of "
        "concerns, comprehensive type annotations, graceful error degradation, and "
        "user-centred design — all hallmarks of professional cybersecurity tooling.",
        "success")

    doc.add_page_break()

    # ================================================================
    # APPENDIX A — FILE STRUCTURE
    # ================================================================
    add_heading(doc, "Appendix A: File Structure", 1)
    add_code_block(doc, """\
Email-Header-Phishing-Investigation-Tool/
│
├── email_forensic_analyzer.py   (1,627 lines)  # Backend analysis engine
├── forensic_gui.py              (2,214 lines)  # CustomTkinter GUI
├── requirements.txt                            # Python package dependencies
├── README.md                                   # Basic execution instructions
├── .gitignore                                  # Git exclusion rules
│
├── email.eml                                   # Sample phishing email (534 lines)
└── sample.eml                                  # Minimal test email (61 lines)""")

    add_heading(doc, "Project Statistics", 2)
    add_section_table(doc,
        ["Metric", "Value"],
        [
            ["Total Python Lines of Code",        "3,841"],
            ["Backend Module (analyzer.py)",      "1,627 lines"],
            ["GUI Module (forensic_gui.py)",       "2,214 lines"],
            ["Major Backend Methods",              "23"],
            ["Regex Patterns Defined",             "50+"],
            ["Analysis Phases",                    "8–10"],
            ["GUI Tabs",                           "10"],
            ["Supported Risky Extensions",         "48"],
            ["Third-party Dependencies",           "5"],
            ["External API Integrations",          "3"],
            ["Git Commits",                        "8"],
            ["Sample Email Files",                 "2"],
        ])

    # ================================================================
    # APPENDIX B — RISK SCORE WEIGHT REFERENCE
    # ================================================================
    add_heading(doc, "Appendix B: Risk Score Weight Reference", 1)
    add_section_table(doc,
        ["Rank", "Signal", "Max Points", "% of Total"],
        [
            ["1", "Authentication Failure (SPF/DKIM fail)", "25", "25%"],
            ["2", "URL Display/Href Mismatch",              "20", "20%"],
            ["3", "Risky Attachment Extension",             "15", "15%"],
            ["4", "Young Domain (< 30 days)",               "15", "15%"],
            ["5", "Abusive IP (AbuseIPDB ≥ 25%)",           "10", "10%"],
            ["6", "Urgency Language Patterns",              "5",  "5%"],
            ["7", "Credential Harvesting Language",         "5",  "5%"],
            ["8", "Brand Impersonation Patterns",           "5",  "5%"],
            ["",  "TOTAL",                                  "100","100%"],
        ])

    # ================================================================
    # APPENDIX C — RISKY FILE EXTENSIONS
    # ================================================================
    add_heading(doc, "Appendix C: Supported Risky File Extensions", 1)
    add_section_table(doc,
        ["Category", "Extensions"],
        [
            ["Windows Executables",  ".exe, .scr, .com, .pif, .bat, .cmd"],
            ["PowerShell / Scripts", ".ps1, .psm1, .psd1, .hta, .wsf, .ws, .wsh"],
            ["Web Scripts",          ".js, .jse, .vbs, .vbe"],
            ["Registry / Config",    ".reg, .inf"],
            ["Installers",           ".msi, .msp"],
            ["Office Macros",        ".docm, .xlsm, .pptm, .dotm, .xlam, .xltm"],
            ["Disk / VM Images",     ".iso, .img, .vhd, .vmdk"],
            ["Archives",             ".zip, .rar, .7z, .cab, .tar, .gz, .bz2, .ace"],
            ["Shortcut / LNK",       ".lnk, .url"],
            ["Java",                 ".jar, .class"],
            ["Compiled HTML",        ".chm"],
            ["Other",                ".dll, .sys, .drv, .cpl, .ocx"],
        ])

    # ================================================================
    # APPENDIX D — PHISHING PATTERN CATEGORIES
    # ================================================================
    add_heading(doc, "Appendix D: Phishing Pattern Categories", 1)
    add_section_table(doc,
        ["Category", "Example Trigger Phrases"],
        [
            ["Urgency & Pressure",     "urgent, act now, immediate action required, account suspended, verify your, click immediately, limited time, expires soon"],
            ["Credential Harvesting",  "enter your password, login credentials, credit card, bank account, social security, submit your information, confirm identity"],
            ["Brand Impersonation",    "PayPal, Apple, Microsoft, Amazon, Netflix, Bank of America, Chase, IRS, USPS, FedEx, DHL, Wells Fargo, Google"],
        ])

    # ================================================================
    # FOOTER
    # ================================================================
    add_divider(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"Email Header Phishing Investigation Tool — Project Report  |  "
        f"Confidential  |  Generated {datetime.date(2026, 4, 16).strftime('%B %d, %Y')}  |  "
        f"Author: Maryam Inam"
    )
    r.font.size      = Pt(9)
    r.font.color.rgb = GREY
    r.italic         = True

    doc.save(output_path)
    print(f"Report saved to: {output_path}")


if __name__ == "__main__":
    build_report("e:/Email-Header-Phishing-Investigation-Tool/Project_Report_Email_Phishing_Tool.docx")
